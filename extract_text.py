from __future__ import annotations
import os, re, io, json, math, csv, itertools, mimetypes
from pathlib import Path
from typing import Optional, List, Tuple
import PyPDF2
import pdfminer
import docx
import bs4
import pandas
import PIL
import pytesseract
import chardet
import cv2
import tempfile
from PIL import Image
import numpy as np

try:
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
except:
    pass

class TextExtractor:  # Fixed typo in class name
   
    MAX_VIDEO_FRAMES = 50
    DEFAULT_FRAME_INTERVAL = 0.5
    OCR_LANGUAGES = 'rus+eng'
    DEBUG_MODE = False
    
    @classmethod
    def detect_encoding(cls, raw_bytes: bytes) -> str:
        """Detect text encoding from byte data"""
        if chardet is None:
            return 'utf-8'
        try:
            res = chardet.detect(raw_bytes)
            enc = res.get('encoding') or 'utf-8'
            # Handle common aliases
            if enc.lower() == 'ascii':
                return 'utf-8'
            return enc
        except Exception:
            return 'utf-8'

    @classmethod
    def extract_text_generic(cls, path: Path) -> str:
        """Extract text from generic text files"""
        try:
            raw = path.read_bytes()
            enc = cls.detect_encoding(raw)
            return raw.decode(enc, errors='ignore')
        except Exception as e:
            if cls.DEBUG_MODE:
                print(f"Generic extraction error for {path.name}: {e}")
            return ''
    
    @classmethod
    def extract_text_pdf(cls, path: Path) -> str:
        """Extract text from PDF files"""
        text = ''
        
        # Try pdfminer first (better extraction)
        if pdfminer is not None:
            try:
                from pdfminer.high_level import extract_text as pdfminer_extract
                text = pdfminer_extract(str(path)) or ''
                if text.strip():
                    return text
            except Exception as e:
                if cls.DEBUG_MODE:
                    print(f"PDFMiner error: {e}")
        
        # Fallback to PyPDF2
        if PyPDF2 is not None:
            try:
                with open(path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    pages_text = []
                    for page_num, page in enumerate(reader.pages, 1):
                        try:
                            page_text = page.extract_text() or ''
                            if page_text.strip():
                                pages_text.append(page_text)
                        except Exception as e:
                            if cls.DEBUG_MODE:
                                print(f"Error extracting PDF page {page_num}: {e}")
                            continue
                    text = '\n'.join(pages_text)
                return text
            except Exception as e:
                if cls.DEBUG_MODE:
                    print(f"PyPDF2 error: {e}")
                return ''
        return text

    @classmethod
    def extract_text_docx(cls, path: Path) -> str:
        """Extract text from DOCX files"""
        if docx is None:
            return ''
        try:
            from docx import Document
            doc = Document(str(path))
            parts = []
            
            # Extract paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            
            # Extract tables
            for tbl in doc.tables:
                for row in tbl.rows:
                    row_text = ' \t '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        parts.append(row_text)
            
            return '\n'.join(parts)
        except Exception as e:
            if cls.DEBUG_MODE:
                print(f"DOCX extraction error: {e}")
            return ''

    @classmethod
    def extract_text_html(cls, path: Path) -> str:
        """Extract text from HTML files"""
        txt = cls.extract_text_generic(path)
        if bs4 is None:
            return txt
        try:
            from bs4 import BeautifulSoup
            # Try lxml parser first, fallback to html.parser
            try:
                soup = BeautifulSoup(txt, 'lxml')
            except:
                soup = BeautifulSoup(txt, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text(separator=' ', strip=True)
            return re.sub(r'\s+', ' ', text).strip()
        except Exception as e:
            if cls.DEBUG_MODE:
                print(f"HTML extraction error: {e}")
            return txt

    @classmethod
    def extract_text_rtf(cls, path: Path) -> str:
        """Extract text from RTF files (basic extraction)"""
        raw = cls.extract_text_generic(path)
        # Remove RTF control sequences
        raw = re.sub(r'\\[a-zA-Z]+-?\d*\s?', ' ', raw)
        raw = re.sub(r'[{}]', ' ', raw)
        raw = re.sub(r'\\u\d+\?', '', raw)  # Handle Unicode escapes
        return re.sub(r'\s+', ' ', raw).strip()

    @classmethod
    def extract_text_ipynb(cls, path: Path) -> str:
        """Extract text from Jupyter notebooks"""
        try:
            data = json.loads(path.read_text(encoding='utf-8', errors='ignore'))
            parts = []
            for cell in data.get('cells', []):
                src = cell.get('source', [])
                if isinstance(src, list):
                    cell_text = ''.join(src)
                elif isinstance(src, str):
                    cell_text = src
                else:
                    continue
                
                # Add cell type marker
                cell_type = cell.get('cell_type', 'code')
                if cell_type == 'markdown':
                    parts.append(f"[Markdown]: {cell_text}")
                else:
                    parts.append(cell_text)
            
            return '\n'.join(parts)
        except Exception as e:
            if cls.DEBUG_MODE:
                print(f"IPYNB extraction error: {e}")
            return ''

    @classmethod
    def extract_text_xls(cls, path: Path) -> str:
        """Extract text from Excel files"""
        if pandas is None:
            return ''
        try:
            # Read all sheets
            excel_file = pandas.ExcelFile(str(path))
            all_text = []
            
            for sheet_name in excel_file.sheet_names:
                df = pandas.read_excel(str(path), sheet_name=sheet_name, header=None, dtype=str)
                sheet_text = []
                for _, row in df.iterrows():
                    row_text = ' '.join(str(val) for val in row.dropna().tolist() if str(val).strip())
                    if row_text:
                        sheet_text.append(row_text)
                
                if sheet_text:
                    all_text.append(f"[Sheet: {sheet_name}]")
                    all_text.extend(sheet_text)
            
            return '\n'.join(all_text)
        except Exception as e:
            if cls.DEBUG_MODE:
                print(f"Excel extraction error: {e}")
            return ''

    @classmethod
    def extract_text_image(cls, path: Path, preprocess: bool = True) -> str:
        """Extract text from images using OCR"""
        if PIL is None or pytesseract is None:
            return ''
        
        try:
            from PIL import Image, ImageEnhance, ImageFilter
            
            # Open image
            img = Image.open(str(path))
            
            # Preprocess image for better OCR
            if preprocess:
                # Convert to RGB if needed
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Convert to grayscale
                img = img.convert('L')
                
                # Enhance contrast
                enhancer = ImageEnhance.Contrast(img)
                img = enhancer.enhance(2.0)
                
                # Apply sharpening
                img = img.filter(ImageFilter.SHARPEN)
            
            # Perform OCR
            text = pytesseract.image_to_string(img, lang=cls.OCR_LANGUAGES, config='--psm 3')
            return text.strip()
            
        except Exception as e:
            if cls.DEBUG_MODE:
                print(f"OCR Error for {path.name}: {e}")
            return ''
    
    @classmethod
    def extract_text_doc(cls, path: Path) -> str:
        """Extract text from legacy DOC files (basic extraction)"""
        # For better DOC extraction, consider using antiword or catdoc
        raw = cls.extract_text_generic(path)
        # Remove binary garbage
        raw = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', raw)
        return re.sub(r'\s+', ' ', raw).strip()

    @classmethod
    def extract_text_video(cls, path: Path, frame_interval_sec: Optional[float] = None, 
                          max_frames: Optional[int] = None) -> str:
        """Extract text from video frames using OCR"""
        if cv2 is None:
            print(f"OpenCV (cv2) is not installed. Install with: pip install opencv-python")
            return ''
        
        frame_interval = frame_interval_sec or cls.DEFAULT_FRAME_INTERVAL
        max_frames = max_frames or cls.MAX_VIDEO_FRAMES
        
        cap = None
        all_texts = []
        processed_texts = set()  # Avoid duplicates
        
        try:
            cap = cv2.VideoCapture(str(path))
            if not cap.isOpened():
                print(f"Cannot open video: {path.name}")
                return ''
            
            fps = cap.get(cv2.CAP_PROP_FPS)
            if fps <= 0:
                fps = 30  # Default FPS
            
            total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            duration = total_frames / fps if fps > 0 else 0
            
            print(f"Processing video {path.name}: {duration:.1f} sec, {total_frames} frames, {fps} fps")
            
            frame_step = max(1, int(fps * frame_interval))
            frame_number = 0
            processed_frames = 0
            
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)
                
                while processed_frames < max_frames and frame_number < total_frames:
                    ret, frame = cap.read()
                    if not ret:
                        break
                    
                    if frame_number % frame_step == 0:
                        timestamp = frame_number / fps
                        
                        # Save frame as image
                        temp_image_path = temp_path / f"frame_{frame_number}.jpg"
                        cv2.imwrite(str(temp_image_path), frame)
                        
                        # Extract text from frame
                        text = cls.extract_text_image(temp_image_path, preprocess=True)
                        
                        # Clean up
                        try:
                            temp_image_path.unlink()
                        except:
                            pass
                        
                        # Add meaningful text (avoid duplicates and very short text)
                        if text and len(text.strip()) > 10:
                            text_hash = text.strip()[:100]  # Simple deduplication
                            if text_hash not in processed_texts:
                                processed_texts.add(text_hash)
                                all_texts.append(f"[{timestamp:.1f}s] {text.strip()}")
                                if cls.DEBUG_MODE:
                                    print(f"Found text at {timestamp:.1f}s: {text[:50]}...")
                        
                        processed_frames += 1
                    
                    frame_number += 1
                    
                    # Safety limit
                    if frame_number > 10000:
                        print(f"Reached frame limit (10000) for {path.name}")
                        break
            
            cap.release()
            
            if all_texts:
                result = '\n'.join(all_texts)
                print(f"Extracted {len(result)} chars from {len(all_texts)} frames of {path.name}")
                return result
            else:
                print(f"No text found in {path.name} (processed {processed_frames} frames)")
                if cls.DEBUG_MODE:
                    cls.debug_save_frames(path, frame_step, max_frames=3)
                return ''
                
        except Exception as e:
            print(f"Error processing video {path.name}: {e}")
            if cls.DEBUG_MODE:
                import traceback
                traceback.print_exc()
            return ''
        finally:
            if cap:
                cap.release()

    @classmethod
    def debug_save_frames(cls, path: Path, frame_step: int, max_frames: int = 5):
        """Save frames for debugging OCR issues"""
        try:
            debug_dir = Path('debug_frames')
            debug_dir.mkdir(exist_ok=True)
            
            cap = cv2.VideoCapture(str(path))
            frame_number = 0
            saved = 0
            
            while saved < max_frames:
                ret, frame = cap.read()
                if not ret:
                    break
                
                if frame_number % frame_step == 0:
                    frame_path = debug_dir / f"{path.stem}_frame_{saved}.jpg"
                    cv2.imwrite(str(frame_path), frame)
                    print(f"  → Saved frame {saved+1} for debugging: {frame_path}")
                    saved += 1
                
                frame_number += 1
            
            cap.release()
            print(f"  → Saved {saved} frames to {debug_dir}/")
            print(f"  → Check these frames manually. If text is visible, OCR configuration needs adjustment")
        except Exception as e:
            print(f"Error saving frames: {e}")

    @classmethod
    def extract_text_parquet(cls, path: Path) -> str:
        """Extract text from Parquet files"""
        try:
            import pandas as pd
            df = pd.read_parquet(str(path))
            # Convert to string representation
            return df.to_string(max_rows=100, max_cols=20)
        except ImportError:
            return "Parquet support requires pandas and pyarrow"
        except Exception as e:
            if cls.DEBUG_MODE:
                print(f"Parquet extraction error: {e}")
            return ''
        
    @classmethod
    def extract_text_markdown(cls, path: Path) -> str:
        """Extract clean text from Markdown files"""
        text = cls.extract_text_generic(path)
        # Remove headers
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        # Convert links to just text
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        # Remove images
        text = re.sub(r'!\[[^\]]*\]\([^\)]+\)', '', text)
        # Remove code blocks
        text = re.sub(r'```[\s\S]*?```', '', text)
        # Remove inline code
        text = re.sub(r'`([^`]+)`', r'\1', text)
        # Remove bold/italic markers
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'__([^_]+)__', r'\1', text)
        text = re.sub(r'_([^_]+)_', r'\1', text)
        # Clean up extra whitespace
        return re.sub(r'\n\s*\n', '\n\n', text).strip()
    
    @classmethod
    def extract_text_tiff(cls, path: Path) -> str:
        """Extract text from multi-page TIFF images"""
        if PIL is None or pytesseract is None:
            return ''
        try:
            from PIL import Image
            img = Image.open(str(path))
            text_parts = []
            page_num = 1
            
            try:
                while True:
                    text = pytesseract.image_to_string(img, lang=cls.OCR_LANGUAGES)
                    if text.strip():
                        text_parts.append(f"[Page {page_num}]\n{text.strip()}")
                    page_num += 1
                    img.seek(img.tell() + 1)
            except EOFError:
                pass
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            if cls.DEBUG_MODE:
                print(f"TIFF extraction error: {e}")
            return ''

    @classmethod
    def extract_text_by_extension(cls, path: Path) -> str:
        """Route file to appropriate extraction method based on extension"""
        ext = path.suffix.lower().lstrip('.')
        
        # Video formats
        if ext in {'mp4', 'avi', 'mov', 'mkv', 'wmv', 'flv', 'webm'}:
            return cls.extract_text_video(path)
        
        elif ext == 'pdf':
            return cls.extract_text_pdf(path)
        elif ext == 'docx':
            return cls.extract_text_docx(path)
        elif ext == 'doc':
            return cls.extract_text_doc(path)
        elif ext == 'rtf':
            return cls.extract_text_rtf(path)
        
        elif ext in {'html', 'htm', 'php', 'asp', 'aspx'}:
            return cls.extract_text_html(path)
        
        elif ext == 'ipynb':
            return cls.extract_text_ipynb(path)
        
        elif ext in {'xls', 'xlsx', 'xlsm'}:
            return cls.extract_text_xls(path)
        
        elif ext == 'parquet':
            return cls.extract_text_parquet(path)
        elif ext == 'md':
            return cls.extract_text_markdown(path)
        
        elif ext in {'jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp'}:
            return cls.extract_text_image(path)
        elif ext in {'tiff', 'tif'}:
            return cls.extract_text_tiff(path)
        
        else:
            return cls.extract_text_generic(path)
    
    @classmethod
    def extract_text(cls, path: Path) -> str:
        """Main extraction method with error handling"""
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        try:
            text = cls.extract_text_by_extension(path)
            # Clean up the extracted text
            if text:
                # Remove excessive whitespace but preserve structure
                text = re.sub(r'[ \t]+', ' ', text)
                text = re.sub(r'\n[ \t]+', '\n', text)
                text = re.sub(r'[ \t]+\n', '\n', text)
                return text.strip()
            return ''
        except Exception as e:
            print(f"Extraction failed for {path.name}: {e}")
            if cls.DEBUG_MODE:
                import traceback
                traceback.print_exc()
            return ''
    
    @classmethod
    def extract_text_batch(cls, paths: List[Path]) -> dict:
        """Extract text from multiple files"""
        results = {}
        for path in paths:
            print(f"Processing: {path.name}")
            results[str(path)] = cls.extract_text(path)
        return results